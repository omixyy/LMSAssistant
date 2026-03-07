import os
import logging
import re
from typing import Tuple, Dict, List, Optional
from datetime import datetime
import pypdf
from docx import Document
import base64

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileProcessor:
    """Обработчик файлов различных форматов"""

    SUPPORTED_EXTENSIONS = {
        '.pdf': 'PDF документ',
        '.docx': 'Word документ',
        '.doc': 'Word документ (старый формат)',
        '.txt': 'Текстовый файл',
        '.rtf': 'Rich Text Format'
    }

    def __init__(self, max_file_size_mb: int = 10, extract_images: bool = True):
        self.max_file_size = max_file_size_mb * 1024 * 1024
        self.extract_images = extract_images

    def process(self, file_path: str) -> Tuple[str, Dict]:
        """
        Основной метод обработки файла
        Возвращает (текст, метаданные) - БЕЗ предопределения типа
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f'Файл не найден: {file_path}')

        # Проверка размера
        file_size = os.path.getsize(file_path)
        if file_size > self.max_file_size:
            raise ValueError(
                f'Файл слишком большой: '
                f'{file_size / 1024 / 1024:.1f} МБ '
                f'(макс. {self.max_file_size / 1024 / 1024:.0f} МБ)',
            )

        # Получение расширения
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f'Неподдерживаемый формат: '
                f'{file_ext}. Поддерживаются: '
                f'{list(self.SUPPORTED_EXTENSIONS.keys())}',
            )

        # Базовые метаданные файла (без определения типа!)
        metadata = {
            'filename': os.path.basename(file_path),
            'extension': file_ext,
            'size_bytes': file_size,
            'size_mb': round(file_size / 1024 / 1024, 2),
            'created': datetime.fromtimestamp(
                os.path.getctime(file_path)
            ).isoformat() if hasattr(os, 'getctime') else None,
            'modified': datetime.fromtimestamp(
                os.path.getmtime(file_path)
            ).isoformat()
        }

        # Извлечение текста и изображений
        if file_ext == '.pdf':
            text, images = self._extract_from_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            text, images = self._extract_from_docx(file_path)
        else:
            text = self._extract_text_simple(file_path, file_ext)
            images = []

        # Очистка текста
        text = self._clean_text(text)

        # Статистика (только объективные метрики)
        metadata.update({
            'char_count': len(text),
            'word_count': len(text.split()),
            'images_count': len(images),
            'images': images if self.extract_images else [],
            # НЕТ определения типа - это сделает модель!
        })

        logger.info(f'Файл {file_path} обработан: {metadata["word_count"]} слов, {len(images)} изображений')

        return text, metadata

    def _extract_from_pdf(self, file_path: str) -> Tuple[str, List[Dict]]:
        """Извлечение текста и изображений из PDF"""
        text = ''
        images = []

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages, 1):
                    # Извлечение текста
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f'--- Страница {page_num} ---\n{page_text}\n\n'

                    # Извлечение изображений (опционально)
                    if self.extract_images and hasattr(page, 'images'):
                        for img_idx, img in enumerate(page.images):
                            try:
                                img_base64 = base64.b64encode(img.data).decode('utf-8')
                                images.append({
                                    'page': page_num,
                                    'index': img_idx,
                                    'format': img.name.split('.')[-1] if '.' in img.name else 'unknown',
                                    'size': len(img.data),
                                    'base64': img_base64[:100] + '...'  # Обрезаем для метаданных
                                })
                            except Exception as e:
                                logger.warning(f'Ошибка при извлечении изображения: {e}')

            return text, images

        except Exception as e:
            logger.error(f'Ошибка чтения PDF: {str(e)}')
            raise

    # ... остальные методы остаются без изменений ...
    def _extract_from_docx(self, file_path: str) -> Tuple[str, List[Dict]]:
        """Извлечение текста из DOCX"""
        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n'.join(text_parts), []  # DOCX изображения сложнее извлекать

        except Exception as e:
            logger.error(f'Ошибка чтения DOCX: {str(e)}')
            raise

    def _extract_text_simple(self, file_path: str, file_ext: str) -> str:
        """Простое извлечение текста для TXT и RTF"""
        if file_ext == '.txt':
            return self._extract_from_txt(file_path)
        elif file_ext == '.rtf':
            return self._extract_from_rtf(file_path)
        return ''

    @staticmethod
    def _extract_from_txt(file_path: str) -> str:
        """Извлечение текста из TXT"""
        encodings = ['utf-8', 'cp1251', 'koi8-r', 'latin-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue

        raise ValueError(f'Не удалось прочитать файл {file_path} ни в одной из кодировок: {encodings}')

    @staticmethod
    def _extract_from_rtf(file_path: str) -> str:
        """Простое извлечение текста из RTF"""
        try:
            with open(file_path, 'r', encoding='cp1251', errors='ignore') as file:
                content = file.read()
                import re
                text = re.sub(r'\\[a-z]+[\d]*', ' ', content)
                text = re.sub(r'[{}]', '', text)
                return text
        except Exception as e:
            logger.error(f'Ошибка чтения RTF: {str(e)}')
            raise

    @staticmethod
    def _clean_text(text: str) -> str:
        """Очистка и нормализация текста"""
        if not text:
            return ''

        import re
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s.,!?;:""''()\[\]{}«»—–-]', '', text)
        return text.strip()