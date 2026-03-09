import logging
from typing import Dict, List, Any, Optional
from pathlib import Path
import re

from .base_parser import Parser
from .non_text_elements import Formula, Table

try:
    from docx import Document
    from docx.table import Table as DocxTable

    DOCX_AVAILABLE = True

except ImportError:
    DOCX_AVAILABLE = False
    Document = None

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class DOCXParser(Parser):
    """
    Парсер для DOCX документов
    """

    def __init__(self, config: Optional[Dict] = None):
        super().__init__(config)

        if not DOCX_AVAILABLE:
            logger.warning("python-docx не установлен. DOCXParser не будет работать.")

    @property
    def supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']

    @property
    def parser_name(self) -> str:
        return "DOCXParser"

    def parse(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        Парсинг DOCX документа
        """
        self._log_parse_start(file_path)

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx не установлен. Установите: pip install python-docx")

        result = self._init_result_dict(Path(file_path).name)

        doc = Document(file_path)
        result['metadata']['pages'] = 1  # DOCX не имеет страниц в том же смысле

        text_parts = []
        formulas = []
        tables = []

        # Парсинг параграфов
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            # Проверка на формулы (упрощённо)
            if self.detect_formulas and self._looks_like_formula(text):
                formula = Formula(
                    page=1,
                    text=text,
                    latex=self._text_to_latex(text),
                    bbox=(0, 0, 0, 0),  # DOCX не даёт координат
                    context='',
                    confidence=0.8,
                    formula_type='inline'
                )
                formulas.append(formula)
                text_parts.append(f'«ФОРМУЛА inline: ${formula.latex}$»')
            else:
                text_parts.append(text)

        # Парсинг таблиц
        if self.detect_tables:
            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_table_data(table)
                if table_data:
                    extracted_table = Table(
                        page=1,
                        bbox=(0, 0, 0, 0),
                        data=table_data,
                        caption=f"Таблица {table_idx + 1}"
                    )
                    tables.append(extracted_table)

        result['text'] = '\n'.join(text_parts)
        result['formulas'] = formulas
        result['tables'] = tables
        result['metadata']['formulas_count'] = len(formulas)
        result['metadata']['tables_count'] = len(tables)

        self._log_parse_complete(result)
        return result

    @staticmethod
    def _extract_table_data(table) -> List[List[str]]:
        """Извлечение данных из DOCX таблицы"""
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Извлечение текста из ячейки (может содержать несколько параграфов)
                cell_text = '\n'.join(p.text for p in cell.paragraphs).strip()
                row_data.append(cell_text)
            if any(row_data):  # Пропускаем пустые строки
                data.append(row_data)
        return data

    @staticmethod
    def _looks_like_formula(text: str) -> bool:
        """Проверка, похож ли текст на формулу"""
        # Простые эвристики
        math_patterns = [
            r'[=≠<>]',
            r'\\[a-z]+',
            r'\^[_0-9]',
            r'_{[0-9]}',
            r'\d+\s*[/×]\s*\d+',
            r'[∑∏∫∂∇]',
        ]

        math_score = 0
        for pattern in math_patterns:
            if re.search(pattern, text):
                math_score += 1

        return math_score >= 2

    @staticmethod
    def _text_to_latex(text: str) -> str:
        """Конвертация текста в LaTeX (упрощённо)"""
        latex = text
        replacements = {
            '≠': r'\neq',
            '≤': r'\leq',
            '≥': r'\geq',
            '±': r'\pm',
            '×': r'\times',
            '÷': r'\div',
            '√': r'\sqrt{}',
            '∑': r'\sum',
            '∏': r'\prod',
            '∫': r'\int',
        }
        for orig, latex_eq in replacements.items():
            latex = latex.replace(orig, latex_eq)

        return latex
